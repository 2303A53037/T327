from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import pandas as pd
import os
from io import BytesIO
from docx import Document
from scholarly import scholarly
import sqlite3
import jwt
import datetime
from functools import wraps
import requests
from bs4 import BeautifulSoup

app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = 'your-secret-key'
uploaded_data = None  # To store Excel data temporarily

# --------------------- JWT AUTH ---------------------
def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('x-access-token')
        if not token:
            return jsonify({'error': 'Token is missing!'}), 403
        try:
            jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
        except:
            return jsonify({'error': 'Invalid token!'}), 403
        return f(*args, **kwargs)
    return decorated

@app.route('/login', methods=['POST'])
def login():
    auth = request.get_json()
    if auth['username'] == 'admin' and auth['password'] == 'admin':
        token = jwt.encode({
            'user': auth['username'],
            'exp': datetime.datetime.utcnow() + datetime.timedelta(hours=1)
        }, app.config['SECRET_KEY'], algorithm='HS256')
        return jsonify({'token': token})
    return jsonify({'error': 'Invalid credentials'}), 401

# Initialize Database
def init_db():
    conn = sqlite3.connect('publications.db')
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS publications 
                (id INTEGER PRIMARY KEY, faculty_name TEXT, title TEXT, journal TEXT, year INTEGER)''')
    conn.commit()
    conn.close()

init_db()

# Upload an Excel File
@app.route('/upload', methods=['POST'])
@token_required
def upload_file():
    global uploaded_data
    file = request.files['file']

    if file:
        uploaded_data = pd.read_excel(file)
        uploaded_data['Year'] = pd.to_numeric(uploaded_data['Year'], errors='coerce')
        uploaded_data = uploaded_data.dropna(subset=['Year'])
        uploaded_data['Year'] = uploaded_data['Year'].astype(int)
        return jsonify({"message": "File uploaded successfully!"})
    return jsonify({"error": "No file uploaded"}), 400

# Filter Publications by Year
@app.route('/publications', methods=['GET'])
@token_required
def get_publications():
    if uploaded_data is None:
        return jsonify({"error": "No data uploaded"}), 400

    from_year = int(request.args.get('from'))
    to_year = int(request.args.get('to'))

    filtered_data = uploaded_data[
        (uploaded_data['Year'] >= from_year) & 
        (uploaded_data['Year'] <= to_year)
    ]

    return jsonify(filtered_data.to_dict(orient="records"))

# Export to Excel
@app.route('/export/excel', methods=['GET'])
@token_required
def export_excel():
    global uploaded_data
    if uploaded_data is None:
        return jsonify({"error": "No data uploaded"}), 400

    from_year = request.args.get('from', type=int)
    to_year = request.args.get('to', type=int)

    if from_year is not None and to_year is not None:
        filtered_data = uploaded_data[
            (uploaded_data['Year'] >= from_year) & 
            (uploaded_data['Year'] <= to_year)
        ]
    else:
        filtered_data = uploaded_data

    if filtered_data.empty:
        return jsonify({"error": "No data found for the selected year range"}), 404

    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        journal_pubs = filtered_data[filtered_data['Type'] == 'Journal']
        journal_pubs.to_excel(writer, sheet_name='Journal_Publications', index=False)

        conf_pubs = filtered_data[filtered_data['Type'] == 'Conference']
        conf_pubs.to_excel(writer, sheet_name='Conference_Publications', index=False)

        filtered_data.to_excel(writer, sheet_name='All_Publications', index=False)

    output.seek(0)
    return send_file(
        output,
        download_name="filtered_publications.xlsx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

# Export to Word
@app.route('/export/word', methods=['GET'])
@token_required
def export_word():
    if uploaded_data is None:
        return jsonify({"error": "No data uploaded"}), 400

    from_year = request.args.get('from', type=int)
    to_year = request.args.get('to', type=int)

    if from_year is not None and to_year is not None:
        filtered_data = uploaded_data[
            (uploaded_data['Year'] >= from_year) & 
            (uploaded_data['Year'] <= to_year)
        ]
    else:
        filtered_data = uploaded_data

    doc = Document()
    doc.add_heading("Publication Summary", level=1)

    if from_year and to_year:
        doc.add_paragraph(f"Showing publications from {from_year} to {to_year}")

    for pub_type in ['Journal', 'Conference']:
        type_data = filtered_data[filtered_data['Type'] == pub_type]
        if not type_data.empty:
            doc.add_heading(f"{pub_type} Publications", level=2)
            table = doc.add_table(rows=1, cols=len(type_data.columns))
            hdr_cells = table.rows[0].cells
            for i, col_name in enumerate(type_data.columns):
                hdr_cells[i].text = str(col_name)
            for _, row in type_data.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value)

    output = BytesIO()
    doc.save(output)
    output.seek(0)

    return send_file(
        output,
        download_name="publication_summary.docx",
        as_attachment=True,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# Search Google Scholar
@app.route('/search/scholar', methods=['GET'])
@token_required
def search_scholar():
    faculty_name = request.args.get('name')
    if not faculty_name:
        return jsonify({"error": "No name provided"}), 400

    try:
        search_query = scholarly.search_author(faculty_name)
        author = next(search_query, None)

        if not author:
            return jsonify({"error": "No author found"}), 404

        scholarly.fill(author)
        publications = [
            {
                "Title": pub['bib']['title'],
                "Year": pub['bib'].get('pub_year', 'Unknown'),
                "Type": "Journal" if 'journal' in pub['bib'] else "Conference"
            }
            for pub in author['publications']
        ]

        return jsonify(publications)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Search DBLP
@app.route('/search/dblp', methods=['GET'])
@token_required
def search_dblp():
    faculty_name = request.args.get('name')
    if not faculty_name:
        return jsonify({"error": "No name provided"}), 400

    try:
        query = faculty_name.replace(" ", "+")
        url = f"https://dblp.org/search?q={query}"

        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')

        results = []
        for item in soup.find_all("li", class_="entry"):
            title = item.find("span", class_="title")
            venue = item.find("span", class_="venue")
            year = item.find("span", class_="year")

            if title:
                pub = {
                    "Title": title.text,
                    "Type": "Conference" if venue and ("conf" in venue.text.lower() or "proc" in venue.text.lower()) else "Journal",
                    "Year": year.text if year else "Unknown"
                }
                results.append(pub)

        return jsonify(results)
    except Exception as e:
        return jsonify({"error": str(e)}), 500

# Save to Database
@app.route('/save', methods=['POST'])
@token_required
def save_to_db():
    if uploaded_data is None:
        return jsonify({"error": "No data uploaded"}), 400

    try:
        conn = sqlite3.connect('publications.db')
        c = conn.cursor()

        for _, row in uploaded_data.iterrows():
            c.execute(
                "INSERT INTO publications (faculty_name, title, journal, year) VALUES (?, ?, ?, ?)", 
                (row['Faculty Name'], row['Title'], row['Journal'], row['Year'])
            )

        conn.commit()
        conn.close()

        return jsonify({"message": "Data saved successfully!"})
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True)
